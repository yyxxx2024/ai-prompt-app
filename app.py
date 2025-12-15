import streamlit as st
from openai import OpenAI
import base64
import json
import requests
import time
import io
import hashlib
import qrcode
from io import BytesIO
import re
from datetime import datetime

# 尝试导入 docx
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ================= 1. 页面基本设置 =================
st.set_page_config(page_title="AI 提示词魔法师 Pro", page_icon="🍊", layout="centered")

# ================= 🎨 UI 样式 =================
def add_modern_light_style():
    st.markdown("""
    <style>
        .stApp { background-color: #f7f9fb; font-family: 'PingFang SC', sans-serif; color: #2c3e50; }
        .block-container { background-color: #ffffff; border-radius: 12px; padding: 3rem 2rem; box-shadow: 0 4px 20px rgba(0,0,0,0.06); max-width: 850px; }
        h1 { color: #1a1a1a; font-weight: 700 !important; letter-spacing: -1px; }
        .stTextArea textarea, .stTextInput input, .stSelectbox div[data-baseweb="select"] > div {
            background-color: #f8f9fa !important; border: 1px solid #e9ecef !important; border-radius: 8px; color: #495057 !important;
        }
        .stTextArea textarea:focus, .stTextInput input:focus { border-color: #ff8c42 !important; box-shadow: 0 0 0 2px rgba(255, 140, 66, 0.2); }
        div.stButton > button[kind="primary"] {
            background: linear-gradient(135deg, #ff9a44 0%, #fc6076 100%); border: none; color: white !important;
            padding: 12px 28px; border-radius: 8px; font-weight: 600; font-size: 16px; width: 100%; transition: all 0.3s;
            box-shadow: 0 4px 12px rgba(252, 96, 118, 0.3);
        }
        div.stButton > button[kind="primary"]:hover { transform: translateY(-2px); box-shadow: 0 6px 16px rgba(252, 96, 118, 0.4); }
        div.stButton > button[kind="secondary"] { border: 1px solid #eee; color: #666; background-color: white; border-radius: 6px; }
        div.stButton > button[kind="secondary"]:hover { border-color: #ff9a44; color: #ff9a44; background-color: #fff8f0; }
        .stTabs [data-baseweb="tab-list"] { gap: 24px; }
        .stTabs [aria-selected="true"] { color: #fc6076 !important; border-bottom-color: #fc6076 !important; }
        .stCode { background-color: #f8f9fa !important; border: 1px solid #eee; border-radius: 8px; }
        .stInfo { background-color: #f0f7ff; color: #0052cc; border: none; border-radius: 8px; }
        .stSuccess { background-color: #e6fffa; color: #009975; border: none; border-radius: 8px; }
    </style>
    """, unsafe_allow_html=True)

add_modern_light_style()

# ================= 🛠️ Gitee 核心函数 =================
def get_gitee_config():
    return {
        "token": st.secrets.get("GITEE_TOKEN", ""),
        "owner": st.secrets.get("GITEE_OWNER", ""),
        "repo": st.secrets.get("GITEE_REPO", ""),
        "branch": st.secrets.get("GITEE_BRANCH", "main")
    }

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def validate_username(username):
    return bool(re.match(r"^[a-zA-Z0-9_-]{3,20}$", username))

def get_all_users():
    try:
        cfg = get_gitee_config()
        url = f"https://gitee.com/api/v5/repos/{cfg['owner']}/{cfg['repo']}/contents/users.json"
        res = requests.get(url, params={"access_token": cfg['token'], "ref": cfg['branch']})
       
        if res.status_code == 200:
            content = res.json()['content']
            decoded = base64.b64decode(content).decode('utf-8')
            data = json.loads(decoded)
            if isinstance(data, list):
                return {}
            return data
        elif res.status_code == 404:
            return {}
        return {}
    except:
        return {}

def register_new_user(username, password):
    if not validate_username(username):
        return False, "❌ 用户名只能包含字母数字下划线和连字符，长度3-20"
   
    users = get_all_users()
    if isinstance(users, list): users = {}
    if username in users:
        return False, "❌ 用户名已存在"
   
    users[username] = hash_password(password)
   
    try:
        cfg = get_gitee_config()
        url = f"https://gitee.com/api/v5/repos/{cfg['owner']}/{cfg['repo']}/contents/users.json"
       
        sha = None
        get_res = requests.get(url, params={"access_token": cfg['token'], "ref": cfg['branch']})
        if get_res.status_code == 200:
            sha = get_res.json()['sha']
       
        new_text = json.dumps(users, ensure_ascii=False, indent=4)
        new_b64 = base64.b64encode(new_text.encode('utf-8')).decode('utf-8')
       
        payload = {
            "access_token": cfg['token'],
            "content": new_b64,
            "message": f"Register user {username}",
            "branch": cfg['branch']
        }
        if sha: payload["sha"] = sha
       
        res = requests.put(url, json=payload)
       
        if res.status_code in [200, 201]:
            return True, "✅ 注册成功！已自动登录"
        else:
            return False, f"注册失败: {res.text}"
    except Exception as e:
        return False, str(e)

# --- 数据存储系统 ---
def get_user_filename(username):
    return f"prompts_{username}.json"

def load_data(username):
    try:
        cfg = get_gitee_config()
        if not cfg["token"]: return []
        url = f"https://gitee.com/api/v5/repos/{cfg['owner']}/{cfg['repo']}/contents/{get_user_filename(username)}"
        res = requests.get(url, params={"access_token": cfg['token'], "ref": cfg['branch']})
        if res.status_code == 200:
            return json.loads(base64.b64decode(res.json()['content']).decode('utf-8'))
        return []
    except:
        return []

def save_data_item(new_item, username):
    new_item["timestamp"] = datetime.now().isoformat()
    try:
        cfg = get_gitee_config()
        url = f"https://gitee.com/api/v5/repos/{cfg['owner']}/{cfg['repo']}/contents/{get_user_filename(username)}"
        res = requests.get(url, params={"access_token": cfg['token'], "ref": cfg['branch']})
        sha, data_list = None, []
        if res.status_code == 200:
            sha = res.json()['sha']
            try: data_list = json.loads(base64.b64decode(res.json()['content']).decode('utf-8'))
            except: pass
        data_list.append(new_item)
        new_b64 = base64.b64encode(json.dumps(data_list, ensure_ascii=False, indent=4).encode('utf-8')).decode('utf-8')
        payload = {
            "access_token": cfg['token'],
            "content": new_b64,
            "message": f"Add prompt by {username}",
            "branch": cfg['branch']
        }
        if sha: payload["sha"] = sha
        requests.put(url, json=payload)
        st.toast(f"✅ 已保存到 {username} 的宝库")
        time.sleep(1)
    except Exception as e:
        st.error(f"保存出错: {e}")

def delete_data_item(index, username):
    try:
        cfg = get_gitee_config()
        url = f"https://gitee.com/api/v5/repos/{cfg['owner']}/{cfg['repo']}/contents/{get_user_filename(username)}"
        res = requests.get(url, params={"access_token": cfg['token'], "ref": cfg['branch']})
        if res.status_code == 200:
            info = res.json()
            data = json.loads(base64.b64decode(info['content']).decode('utf-8'))
            if 0 <= index < len(data):
                data.pop(index)
                new_b64 = base64.b64encode(json.dumps(data, ensure_ascii=False, indent=4).encode('utf-8')).decode('utf-8')
                requests.put(url, json={
                    "access_token": cfg['token'],
                    "content": new_b64,
                    "sha": info['sha'],
                    "message": f"Delete by {username}",
                    "branch": cfg['branch']
                })
                st.toast("🗑️ 删除成功")
                time.sleep(1)
                st.rerun()
    except:
        pass

def generate_word(data):
    if not HAS_DOCX: return None
    doc = Document()
    doc.add_heading('🌟 我的 AI 提示词宝库', 0)
    doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"用户名: {st.session_state.current_user}\n")
   
    cats = sorted(list(set([d['category'] for d in data])))
    for cat in cats:
        doc.add_heading(f"📂 {cat}", level=1)
        items = sorted([d for d in data if d['category'] == cat],
                       key=lambda x: x.get('timestamp', ''), reverse=True)
        for item in items:
            doc.add_heading(item.get('desc', '无标题'), level=2)
            doc.add_paragraph(f"时间: {item.get('timestamp', '未知')}")
            doc.add_paragraph(item['prompt'])
            doc.add_paragraph("-" * 50)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def encode_image(file):
    return base64.b64encode(file.getvalue()).decode('utf-8')

def custom_select(label, options, key_suffix):
    selected = st.selectbox(label, ["不指定"] + options + ["📝 自定义输入..."], key=f"sel_{key_suffix}")
    if selected == "📝 自定义输入...":
        val = st.text_input(f"请输入 {label}", key=f"txt_{key_suffix}")
        return val if val.strip() else "不指定"
    return selected

def generate_qr_code(url):
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    return img_byte_arr.getvalue()

# ================= 🚀 初始化 =================
if "current_user" not in st.session_state:
    st.session_state.current_user = None
if "last_results" not in st.session_state:
    st.session_state.last_results = None

# URL 自动登录逻辑（兼容最新 Streamlit）
if not st.session_state.current_user:
    params = st.query_params.to_dict()
    u_arg = params.get("u")
    p_arg = params.get("p")
    if u_arg and p_arg:
        users_db = get_all_users()
        try:
            decoded_p = base64.b64decode(p_arg).decode('utf-8')
            hashed_p = hash_password(decoded_p)
            if u_arg in users_db and users_db[u_arg] == hashed_p:
                st.session_state.current_user = u_arg
                st.toast(f"🎉 扫码登录成功！欢迎 {u_arg}")
                time.sleep(1)
                st.rerun()
        except:
            pass

# ================= 🔐 侧边栏 =================
with st.sidebar:
    st.markdown("### 🔐 账户中心")
   
    if st.session_state.current_user:
        st.success(f"👤 已登录: **{st.session_state.current_user}**")
       
        with st.expander("📱 生成免密二维码"):
            st.caption("朋友扫此码可直接登录你的账号")
            try:
                confirm_pass = st.text_input("验证当前密码以生成", type="password", key="qr_pass")
                if confirm_pass:
                    users_db = get_all_users()
                    if users_db.get(st.session_state.current_user) == hash_password(confirm_pass):
                        b64_pass = base64.b64encode(confirm_pass.encode()).decode()
                        current_url = "https://your-app-name.streamlit.app"  # 请替换为你的实际部署域名
                        login_link = f"{current_url}?u={st.session_state.current_user}&p={b64_pass}"
                        qr_img = generate_qr_code(login_link)
                        st.image(qr_img, caption="微信扫一扫，免密直连")
                        st.code(login_link, language=None)
                    else:
                        st.error("密码错误")
            except:
                pass
       
        if st.button("退出登录"):
            st.session_state.current_user = None
            st.rerun()
    else:
        auth_mode = st.radio("选择模式", ["登录", "注册新账号"], horizontal=True)
        user_input_name = st.text_input("用户名", placeholder="英文/数字/下划线")
        user_input_pass = st.text_input("密码", type="password")
       
        if auth_mode == "登录":
            if st.button("登录", type="primary"):
                if not validate_username(user_input_name):
                    st.error("用户名格式无效")
                else:
                    users_db = get_all_users()
                    hashed_pw = hash_password(user_input_pass)
                    if user_input_name in users_db and users_db[user_input_name] == hashed_pw:
                        st.session_state.current_user = user_input_name
                        st.success("✅ 登录成功！")
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error("❌ 用户名或密码错误")
        else:
            if st.button("✨ 立即注册"):
                success, msg = register_new_user(user_input_name, user_input_pass)
                if success:
                    st.session_state.current_user = user_input_name
                    st.success(msg)
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error(msg)
   
    st.markdown("---")
    st.markdown("### ⚙️ 系统设置")
    base_url = st.text_input("API 地址", value="https://hk-api.gptbest.vip/v1")
    text_model = st.text_input("文本模型", value="deepseek-chat")
    vision_model = st.text_input("视觉模型", value="gpt-4o-mini")
    st.caption("支持所有 OpenAI 兼容格式的 API")

# ================= 🏗️ 主界面 =================
st.markdown("# 🍊 AI Prompt Wizard <small>Pro</small>", unsafe_allow_html=True)

if not st.session_state.current_user:
    st.info("👋 欢迎！请在左侧 **登录** 或 **注册** 一个账号开始使用。")
    st.stop()

tab1, tab2, tab3 = st.tabs(["📝 生成提示词", "🖼️ 图片反推", "🌟 我的云端宝库"])

# --- Tab 1: 生成提示词 ---
with tab1:
    user_input = st.text_area("输入你的创意描述", height=100, label_visibility="collapsed",
                              placeholder="例如：一个极简风格的白色美术馆，夕阳下，混凝土材质...")
    c1, c2 = st.columns(2)
    with c1:
        ratio = st.selectbox("画幅比例", ["--ar 16:9", "--ar 3:4", "--ar 1:1", "--ar 9:16", "--ar 4:3", "--ar 21:9"])
    with c2:
        mode = st.selectbox("生成模式", [
            "🏗️ 建筑效果图 (ArchViz)",
            "📐 建筑设计 (Design Concept)",
            "标准模式 (MJ/SD)",
            "自然语言 (Google Gemini/Flux)",
            "二次元 (Niji)",
            "写实摄影"
        ])

    with st.expander("🎨 高级参数配置 (支持自定义)", expanded=True):
        if "效果图" in mode:
            st.caption("🏗️ **专业建筑效果图参数**")
            ac1, ac2, ac3 = st.columns(3)
            with ac1: viz_view = custom_select("视点", ["人视", "半鸟瞰", "顶视", "虫视", "轴测"], "v1")
            with ac2: viz_time = custom_select("时刻", ["黄金时刻", "蓝调时刻", "正午阳光", "阴天", "雨夜", "夜晚灯光"], "v2")
            with ac3: viz_env = custom_select("环境", ["城市街道", "森林湖畔", "滨水", "雪山", "沙漠", "屋顶花园"], "v3")
            ac4, ac5, ac6 = st.columns(3)
            with ac4: viz_render = custom_select("渲染引擎", ["V-Ray", "Unreal Engine 5", "Lumion", "Corona", "Enscape"], "v4")
            with ac5: viz_mat = custom_select("主要材质", ["混凝土", "玻璃幕墙", "木格栅", "金属", "红砖", "白色涂料"], "v5")
            with ac6: viz_mood = custom_select("氛围", ["史诗壮丽", "宁静禅意", "未来科幻", "极简纯粹", "温暖治愈"], "v6")
        elif "建筑设计" in mode:
            st.caption("📐 **建筑设计概念参数**")
            d1, d2, d3 = st.columns(3)
            with d1: des_type = custom_select("图纸类型", ["手绘草图", "轴测图", "平面图", "剖面图", "白模", "爆炸图"], "d1")
            with d2: des_style = custom_select("建筑师风格", ["扎哈·哈迪德", "勒·柯布西耶", "安藤忠雄", "BIG", "解构主义", "极简主义"], "d2")
            with d3: des_scale = custom_select("建筑尺度", ["摩天大楼", "博物馆", "独栋别墅", "城市综合体", "小型装置"], "d3")
            d4, d5, d6 = st.columns(3)
            with d4: des_mat = custom_select("模型材质", ["卡纸", "椴木", "亚克力", "金属丝", "3D打印"], "d4")
            with d5: des_bg = custom_select("背景", ["纯白背景", "网格纸", "牛皮纸", "阴影投影"], "d5")
            with d6: des_detail = custom_select("细节程度", ["高度详细结构", "概念抽象", "构造节点"], "d6")
        else:
            col_a, col_b, col_c = st.columns(3)
            with col_a: lighting = custom_select("光线", ["柔和自然光", "戏剧性电影光", "霓虹夜景", "黄金小时"], "g1")
            with col_b: camera = custom_select("镜头", ["广角镜头", "50mm标准", "长焦压缩", "鱼眼", "微距"], "g2")
            with col_c: mood = custom_select("氛围", ["梦幻唯美", "史诗震撼", "黑暗阴郁", "赛博朋克", "复古怀旧"], "g3")
       
        st.markdown("---")
        mp1, mp2 = st.columns(2)
        with mp1: stylize = st.slider("风格化强度 (--s)", 0, 1000, 250)
        with mp2: chaos = st.slider("多样性 (--c)", 0, 100, 0)
        negative_prompt = st.text_input("🚫 负面提示 (--no)", value="text, watermark, blurry, low quality, deformed, ugly")

    if st.button("🚀 立即生成双方案", type="primary"):
        try:
            client = OpenAI(api_key=st.secrets["API_KEY"], base_url=base_url)
            details = []
            if "效果图" in mode:
                for val, name in [(viz_view, "视点"), (viz_time, "时刻"), (viz_env, "环境"), (viz_render, "渲染"), (viz_mat, "材质"), (viz_mood, "氛围")]:
                    if val != "不指定": details.append(f"{name}: {val}")
            elif "建筑设计" in mode:
                for val, name in [(des_type, "类型"), (des_style, "风格"), (des_scale, "尺度"), (des_mat, "材质"), (des_bg, "背景"), (des_detail, "细节")]:
                    if val != "不指定": details.append(f"{name}: {val}")
            else:
                for val, name in [(lighting, "光线"), (camera, "镜头"), (mood, "氛围")]:
                    if val != "不指定": details.append(f"{name}: {val}")

            sys_msg = """
            你是一个专业的AI提示词工程师。请根据用户输入和参数，生成两个方案：
            - Plan A：忠实于用户描述，精细优化
            - Plan B：更有创意和艺术性突破
            输出格式严格如下（不要多余文字）：
            ===PLAN_A_CN===
            [中文描述]
            ===PLAN_A_EN===
            [英文提示词]
            ===PLAN_B_CN===
            [中文描述]
            ===PLAN_B_EN===
            [英文提示词]
            """

            full_req = f"用户描述: {user_input}\n参数: {', '.join(details) if details else '无'}\n模式: {mode}"

            with st.spinner('AI 思考中，请稍等...'):
                resp = client.chat.completions.create(
                    model=text_model,
                    temperature=0.8,
                    messages=[{"role": "system", "content": sys_msg}, {"role": "user", "content": full_req}]
                )
                raw = resp.choices[0].message.content

                try:
                    parts = raw.split("===PLAN_A_EN===")
                    p1_cn = parts[0].replace("===PLAN_A_CN===", "").strip()
                    rest1 = parts[1]
                   
                    parts2 = rest1.split("===PLAN_B_CN===")
                    p1_en = parts2[0].strip()
                    rest2 = parts2[1]
                   
                    parts3 = rest2.split("===PLAN_B_EN===")
                    p2_cn = parts3[0].strip()
                    p2_en = parts3[1].strip()
                except:
                    st.error("解析失败，已回退显示原始输出")
                    st.code(raw)
                    p1_cn = p1_en = p2_cn = p2_en = raw

                suffix = f" {ratio}"
                if "自然语言" not in mode:
                    suffix += f" --s {stylize} --c {chaos}"
                    if negative_prompt.strip():
                        suffix += f" --no {negative_prompt}"

                st.session_state.last_results = {
                    "p1_cn": p1_cn,
                    "p1_en": p1_en + suffix,
                    "p2_cn": p2_cn,
                    "p2_en": p2_en + suffix
                }
        except Exception as e:
            st.error(f"API 调用失败: {e}")

    # ================ 优化后的结果展示 ================
    if st.session_state.last_results:
        res = st.session_state.last_results
        st.divider()
        
        col_a, col_b = st.columns(2)
        
        with col_a:
            st.markdown("#### 🅰️ 方案 A（忠实版）")
            
            # 中文描述 + 右上角复制按钮
            st.markdown(f"""
            <div style="position: relative;">
                <div style="background-color: #f0f7ff; border-left: 4px solid #0052cc; padding: 14px 18px; border-radius: 8px; margin-bottom: 16px; line-height: 1.6;">
                    {res['p1_cn'].replace(chr(10), '<br>')}
                </div>
                <button onclick="navigator.clipboard.writeText(`{res['p1_cn'].replace('`', '\\`')}`); this.innerText='✓ 已复制'" 
                        style="position: absolute; top: 10px; right: 10px; background: rgba(0,82,204,0.12); border: none; border-radius: 6px; padding: 8px 12px; cursor: pointer; font-size: 13px; color: #0052cc; font-weight: 600;">
                    📋 复制中文
                </button>
            </div>
            """, unsafe_allow_html=True)
            
            # 英文提示词（高度拉长 + 自带复制按钮）
            st.code(res['p1_en'], language="text", height=220)
            
            if st.button("❤️ 收藏方案 A", key="save_a"):
                save_data_item({
                    "category": "生成记录",
                    "desc": res["p1_cn"][:30] + "..." if len(res["p1_cn"]) > 30 else res["p1_cn"],
                    "prompt": res["p1_en"]
                }, st.session_state.current_user)

        with col_b:
            st.markdown("#### 🅱️ 方案 B（创意版）")
            
            # 中文描述 + 右上角复制按钮
            st.markdown(f"""
            <div style="position: relative;">
                <div style="background-color: #f0f7ff; border-left: 4px solid #0052cc; padding: 14px 18px; border-radius: 8px; margin-bottom: 16px; line-height: 1.6;">
                    {res['p2_cn'].replace(chr(10), '<br>')}
                </div>
                <button onclick="navigator.clipboard.writeText(`{res['p2_cn'].replace('`', '\\`')}`); this.innerText='✓ 已复制'" 
                        style="position: absolute; top: 10px; right: 10px; background: rgba(0,82,204,0.12); border: none; border-radius: 6px; padding: 8px 12px; cursor: pointer; font-size: 13px; color: #0052cc; font-weight: 600;">
                    📋 复制中文
                </button>
            </div>
            """, unsafe_allow_html=True)
            
            # 英文提示词（高度拉长 + 自带复制按钮）
            st.code(res['p2_en'], language="text", height=220)
            
            if st.button("❤️ 收藏方案 B", key="save_b"):
                save_data_item({
                    "category": "生成记录",
                    "desc": res["p2_cn"][:30] + "..." if len(res["p2_cn"]) > 30 else res["p2_cn"],
                    "prompt": res["p2_en"]
                }, st.session_state.current_user)

# --- Tab 2: 图片反推 ---
with tab2:
    st.markdown("### 🖼️ 上传图片 → 自动反推高质量 Prompt")
    up_file = st.file_uploader("支持 JPG / PNG", type=["jpg", "jpeg", "png"], label_visibility="collapsed")
    if up_file and st.button("🔍 开始反推", type="primary"):
        try:
            client = OpenAI(api_key=st.secrets["API_KEY"], base_url=base_url)
            b64 = encode_image(up_file)
            with st.spinner('视觉模型分析中...'):
                resp = client.chat.completions.create(
                    model=vision_model,
                    temperature=0.3,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "请仔细观察这张图片，输出高质量提示词。\n格式严格如下：\nCN: [详细中文描述]\nEN: [专业英文 Midjourney / Stable Diffusion Prompt，带必要参数]"},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
                        ]
                    }]
                )
            raw = resp.choices[0].message.content
            if "EN:" in raw and "CN:" in raw:
                cn = raw.split("EN:")[0].replace("CN:", "").strip()
                en = raw.split("EN:")[1].strip()
                st.image(up_file, caption="原图", width=300)
                
                # 中文描述同样加复制按钮
                st.markdown(f"""
                <div style="position: relative;">
                    <div style="background-color: #f0f7ff; border-left: 4px solid #0052cc; padding: 14px 18px; border-radius: 8px; margin: 16px 0;">
                        <strong>中文描述：</strong><br>{cn.replace(chr(10), '<br>')}
                    </div>
                    <button onclick="navigator.clipboard.writeText(`{cn.replace('`', '\\`')}`); this.innerText='✓ 已复制'" 
                            style="position: absolute; top: 10px; right: 10px; background: rgba(0,82,204,0.12); border: none; border-radius: 6px; padding: 8px 12px; cursor: pointer; font-size: 13px; color: #0052cc; font-weight: 600;">
                        📋 复制中文
                    </button>
                </div>
                """, unsafe_allow_html=True)
                
                st.code(en, height=180)
                if st.button("❤️ 收藏此提示词", key="save_reverse"):
                    save_data_item({
                        "category": "图片反推",
                        "desc": cn[:30] + "..." if len(cn) > 30 else cn,
                        "prompt": en
                    }, st.session_state.current_user)
            else:
                st.code(raw)
        except Exception as e:
            st.error(f"反推失败: {e}")

# --- Tab 3: 云端宝库 ---
with tab3:
    curr_user = st.session_state.current_user
    st.markdown(f"#### 🌟 {curr_user} 的私人提示词宝库")
    data = load_data(curr_user)
   
    col_h, col_b = st.columns([3, 1])
    with col_b:
        if data and HAS_DOCX:
            docx = generate_word(data)
            st.download_button(
                "📥 导出为 Word",
                data=docx,
                file_name=f"{curr_user}_AI提示词宝库_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
   
    st.divider()
   
    with st.expander("➕ 手动添加新提示词"):
        with st.form("add_manual"):
            existing_cats = sorted(list(set([d['category'] for d in data]))) if data else []
            c_mode = st.selectbox("分类", ["📝 新建分类"] + existing_cats)
            new_cat = st.text_input("新分类名称") if c_mode == "📝 新建分类" else c_mode
            desc = st.text_input("备注标题（可选）")
            content = st.text_area("提示词内容", height=150)
            submitted = st.form_submit_button("💾 保存到宝库")
            if submitted:
                if not content.strip():
                    st.warning("内容不能为空")
                else:
                    save_data_item({
                        "category": new_cat or "未分类",
                        "desc": desc or "手动添加",
                        "prompt": content.strip()
                    }, curr_user)
                    st.rerun()
   
    if not data:
        st.info("🎉 宝库还是空的！快去生成或反推一些提示词吧～")
    else:
        filter_cat = st.selectbox("🔍 筛选分类", ["全部"] + sorted(list(set([d['category'] for d in data]))))
        filtered = [d for d in data if filter_cat == "全部" or d['category'] == filter_cat]
       
        for i in range(len(filtered)-1, -1, -1):
            item = filtered[i]
            with st.container(border=True):
                col1, col2 = st.columns([8, 1])
                with col1:
                    timestamp = item.get('timestamp', '未知时间')
                    st.markdown(f"**🏷️ {item['category']}** ・ {item.get('desc', '')} ・ 🕙 {timestamp.split('T')[0]}")
                    st.code(item['prompt'], language=None)
                with col2:
                    if st.button("🗑️", key=f"del_{i}_{hash(item['prompt'])}"):
                        real_index = data.index(item)
                        delete_data_item(real_index, curr_user)

st.caption("🍊 AI Prompt Wizard Pro | 基于 Gitee 云存储 | 2025 最新优化版")
